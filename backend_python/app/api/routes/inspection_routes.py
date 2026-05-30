import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from uuid import uuid4

from fastapi import APIRouter, Depends, File, Form, HTTPException, Request, UploadFile
from docx import Document
from docx.shared import Inches
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
from sqlalchemy import select
from sqlalchemy.orm import Session, selectinload
from starlette.concurrency import run_in_threadpool
from starlette.datastructures import UploadFile as StarletteUploadFile

from app.api.dependencies import get_current_user, require_roles
from app.database.db import get_db
from app.database.models import ExportRecord, Inspection, InspectionImage, User
from app.schemas.inspection_schema import (
    ExportEmailResponse,
    InspectionResponse,
    ScanContainerIdResponse,
)
from app.services.ai.container_ocr import ContainerOcrError, detect_container_number
from app.services.notification.email_service import (
    EmailConfigurationError,
    send_email_with_attachment,
)

router = APIRouter(tags=["inspections"])

UPLOAD_ROOT = Path("uploads/inspections")
REPORT_ROOT = Path("reports")
OCR_TEMP_ROOT = Path("tmp/ocr_scans")

PHOTO_LABELS = [
    "Container Number",
    "Front",
    "Rear",
    "Left Side",
    "Right Side",
    "Front Left",
    "Front Right",
    "Rear Left",
    "Rear Right",
    "Ceiling",
    "Floor",
    "Door",
    "Lock",
    "CSC Plate",
]

CONTAINER_NUMBER_PATTERN = re.compile(r"[A-Z]{4}\d{7}")
MAX_IMAGE_SIZE = 10 * 1024 * 1024
ALLOWED_IMAGE_SUFFIXES = {".jpg", ".jpeg", ".png", ".webp"}
ALLOWED_IMAGE_CONTENT_TYPES = {
    "image/jpeg",
    "image/png",
    "image/webp",
    "application/octet-stream",
}
CHUNK_SIZE = 1024 * 1024


def _ensure_storage() -> None:
    UPLOAD_ROOT.mkdir(parents=True, exist_ok=True)
    REPORT_ROOT.mkdir(parents=True, exist_ok=True)
    OCR_TEMP_ROOT.mkdir(parents=True, exist_ok=True)


def _public_url(request: Request, path: str) -> str:
    return str(request.base_url).rstrip("/") + "/" + path.lstrip("/")


def _inspection_to_dict(inspection: Inspection) -> dict[str, Any]:
    images = [
        {
            "angle": image.angle,
            "label": image.label,
            "url": image.url,
        }
        for image in sorted(inspection.images, key=lambda item: item.angle)
    ]
    return {
        "id": inspection.id,
        "container_number": inspection.container_number,
        "booking_number": inspection.booking_number,
        "truck_number": inspection.truck_number,
        "worker_name": inspection.worker_name,
        "port_name": inspection.port_name,
        "notes": inspection.notes or "",
        "status": inspection.status,
        "created_at": inspection.created_at.isoformat(),
        "updated_at": inspection.updated_at.isoformat(),
        "image_urls": [image["url"] for image in images],
        "images": images,
    }


def _find_inspection(db: Session, inspection_id: str) -> Inspection:
    inspection = db.scalar(
        select(Inspection)
        .options(selectinload(Inspection.images))
        .where(Inspection.id == inspection_id)
    )

    if inspection is not None:
        return inspection

    raise HTTPException(status_code=404, detail="Inspection not found")


def _inspection_report_dir(inspection_id: str) -> Path:
    report_dir = REPORT_ROOT / inspection_id
    report_dir.mkdir(parents=True, exist_ok=True)
    return report_dir


def _safe_file_suffix(filename: str | None) -> str:
    suffix = Path(filename or "").suffix.lower()
    if suffix in ALLOWED_IMAGE_SUFFIXES:
        return suffix
    raise HTTPException(
        status_code=400,
        detail=f"Unsupported image file type: {suffix or 'unknown'}",
    )


def _safe_label(label: str) -> str:
    return re.sub(r"[^a-z0-9_]+", "_", label.lower()).strip("_")


def _validate_image_upload(image: UploadFile) -> str:
    suffix = _safe_file_suffix(image.filename)
    content_type = (image.content_type or "").lower()

    if content_type and content_type not in ALLOWED_IMAGE_CONTENT_TYPES:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported image content type: {content_type}",
        )

    return suffix


async def _save_validated_upload(image: UploadFile, file_path: Path) -> None:
    total_size = 0

    with file_path.open("wb") as buffer:
        while chunk := await image.read(CHUNK_SIZE):
            total_size += len(chunk)

            if total_size > MAX_IMAGE_SIZE:
                buffer.close()
                file_path.unlink(missing_ok=True)
                raise HTTPException(
                    status_code=413,
                    detail="Image is too large. Maximum size is 10MB.",
                )

            buffer.write(chunk)

    if total_size == 0:
        file_path.unlink(missing_ok=True)
        raise HTTPException(status_code=400, detail="Uploaded image is empty")


async def _extract_inspection_images(request: Request) -> list[UploadFile]:
    form = await request.form()
    indexed_images: list[UploadFile] = []

    for index in range(len(PHOTO_LABELS)):
        file_obj = form.get(f"image_{index}")

        if not isinstance(file_obj, StarletteUploadFile):
            indexed_images = []
            break

        indexed_images.append(file_obj)

    if len(indexed_images) == len(PHOTO_LABELS):
        return indexed_images

    repeated_images = [
        file_obj
        for file_obj in form.getlist("images")
        if isinstance(file_obj, StarletteUploadFile)
    ]

    if len(repeated_images) == len(PHOTO_LABELS):
        return repeated_images

    bracket_images = [
        file_obj
        for file_obj in form.getlist("images[]")
        if isinstance(file_obj, StarletteUploadFile)
    ]

    if len(bracket_images) == len(PHOTO_LABELS):
        return bracket_images

    raise HTTPException(
        status_code=400,
        detail=(
            f"Expected {len(PHOTO_LABELS)} classified photos using "
            "image_0 through image_13."
        ),
    )


@router.post(
    "/inspections",
    status_code=201,
    response_model=InspectionResponse,
)
async def create_inspection(
    request: Request,
    container_number: str = Form(...),
    booking_number: str = Form(...),
    truck_number: str = Form(...),
    worker_name: str = Form(...),
    port_name: str = Form(...),
    notes: str = Form(""),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_roles("worker", "admin")),
):
    images = await _extract_inspection_images(request)

    if len(images) != len(PHOTO_LABELS):
        raise HTTPException(
            status_code=400,
            detail=f"Expected {len(PHOTO_LABELS)} images, received {len(images)}",
        )

    _ensure_storage()

    inspection_id = str(uuid4())
    inspection_dir = UPLOAD_ROOT / inspection_id
    inspection_dir.mkdir(parents=True, exist_ok=True)

    image_urls: list[str] = []
    image_records: list[dict[str, Any]] = []

    for index, image in enumerate(images):
        suffix = _validate_image_upload(image)
        safe_label = _safe_label(PHOTO_LABELS[index])
        filename = f"{index + 1:02d}_{safe_label}{suffix}"
        file_path = inspection_dir / filename

        await _save_validated_upload(image, file_path)

        relative_path = f"uploads/inspections/{inspection_id}/{filename}"
        image_urls.append(_public_url(request, relative_path))
        image_records.append(
            {
                "angle": index + 1,
                "label": PHOTO_LABELS[index],
                "url": _public_url(request, relative_path),
                "path": str(file_path),
            }
        )

    inspection = Inspection(
        id=inspection_id,
        container_number=container_number.strip().upper(),
        booking_number=booking_number.strip(),
        truck_number=truck_number.strip(),
        worker_name=worker_name.strip() or current_user.full_name,
        port_name=port_name.strip(),
        notes=notes.strip(),
        status="submitted",
        worker_id=current_user.id,
    )
    db.add(inspection)

    for image_record in image_records:
        db.add(
            InspectionImage(
                inspection_id=inspection_id,
                angle=image_record["angle"],
                label=image_record["label"],
                url=image_record["url"],
                path=image_record["path"],
            )
        )

    db.commit()
    db.refresh(inspection)
    inspection = _find_inspection(db, inspection_id)
    return _inspection_to_dict(inspection)


@router.get("/inspections", response_model=list[InspectionResponse])
async def list_inspections(
    status: str | None = None,
    container_number: str | None = None,
    worker_name: str | None = None,
    port_name: str | None = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    query = select(Inspection).options(selectinload(Inspection.images))
    if status:
        query = query.where(Inspection.status == status.lower())
    if container_number:
        query = query.where(Inspection.container_number.ilike(f"%{container_number}%"))
    if worker_name:
        query = query.where(Inspection.worker_name.ilike(f"%{worker_name}%"))
    if port_name:
        query = query.where(Inspection.port_name.ilike(f"%{port_name}%"))
    if current_user.role == "worker":
        query = query.where(Inspection.worker_id == current_user.id)

    query = query.order_by(Inspection.created_at.desc())
    inspections = db.scalars(query).all()
    return [_inspection_to_dict(inspection) for inspection in inspections]


@router.get("/inspections/{inspection_id}", response_model=InspectionResponse)
async def get_inspection(
    inspection_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    inspection = _find_inspection(db, inspection_id)
    if current_user.role == "worker" and inspection.worker_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not allowed")
    return _inspection_to_dict(inspection)


@router.post(
    "/inspections/{inspection_id}/accept",
    response_model=InspectionResponse,
)
async def accept_inspection(
    inspection_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_roles("manager", "admin")),
):
    inspection = _find_inspection(db, inspection_id)
    inspection.status = "accepted"
    inspection.accepted_by_id = current_user.id
    inspection.updated_at = datetime.now(timezone.utc)
    db.commit()
    db.refresh(inspection)
    inspection = _find_inspection(db, inspection_id)
    return _inspection_to_dict(inspection)


@router.post(
    "/inspections/{inspection_id}/reject",
    response_model=InspectionResponse,
)
async def reject_inspection(
    inspection_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_roles("manager", "admin")),
):
    inspection = _find_inspection(db, inspection_id)
    inspection.status = "rejected"
    inspection.rejected_by_id = current_user.id
    inspection.updated_at = datetime.now(timezone.utc)
    db.commit()
    db.refresh(inspection)
    inspection = _find_inspection(db, inspection_id)
    return _inspection_to_dict(inspection)


def _ensure_accepted(inspection: Inspection) -> None:
    if inspection.status != "accepted":
        raise HTTPException(
            status_code=400,
            detail="Only accepted inspections can be exported",
        )


def _email_attachment(
    *,
    inspection: dict[str, Any],
    attachment_path: Path,
    export_label: str,
) -> tuple[bool, str | None, str]:
    subject = f"{export_label} - {inspection['container_number']}"
    body = (
        "Please find the attached container inspection file.\n\n"
        f"Container Number: {inspection['container_number']}\n"
        f"Booking Number: {inspection['booking_number']}\n"
        f"Truck Number: {inspection['truck_number']}\n"
        f"Status: {inspection['status']}\n"
    )

    try:
        result = send_email_with_attachment(
            subject=subject,
            body=body,
            attachment_path=attachment_path,
        )
    except EmailConfigurationError:
        return False, None, "Email settings are not configured yet."
    except Exception as exc:
        return False, None, f"Email sending failed: {exc}"

    return bool(result["sent"]), str(result["to"]), "Email sent successfully."


def _image_path(image: dict[str, Any]) -> Path | None:
    path = image.get("path")
    if not path:
        return None

    image_path = Path(path)
    if image_path.exists() and image_path.suffix.lower() in {".jpg", ".jpeg", ".png"}:
        return image_path

    return None


def _inspection_to_export_dict(inspection: Inspection) -> dict[str, Any]:
    data = _inspection_to_dict(inspection)
    image_paths = {
        image.angle: image.path
        for image in inspection.images
    }

    for image in data["images"]:
        image["path"] = image_paths.get(image["angle"], "")

    return data


def _create_excel_report(inspection: Inspection) -> Path:
    inspection_data = _inspection_to_dict(inspection)
    report_dir = _inspection_report_dir(inspection.id)
    report_path = report_dir / f"inspection_{inspection.id}_excel.xlsx"

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Inspection"
    sheet.column_dimensions["A"].width = 22
    sheet.column_dimensions["B"].width = 44
    sheet.column_dimensions["C"].width = 16
    sheet.column_dimensions["D"].width = 26

    title_fill = PatternFill("solid", fgColor="075DCC")
    sheet["A1"] = "Container Inspection Report"
    sheet["A1"].font = Font(bold=True, color="FFFFFF", size=16)
    sheet["A1"].fill = title_fill
    sheet.merge_cells("A1:D1")

    rows = [
        ("Container Number", inspection_data["container_number"]),
        ("Booking Number", inspection_data["booking_number"]),
        ("Truck Number", inspection_data["truck_number"]),
        ("Worker", inspection_data["worker_name"]),
        ("Port / Location", inspection_data["port_name"]),
        ("Status", inspection_data["status"]),
        ("Submitted At", inspection_data["created_at"]),
        ("Notes", inspection_data.get("notes", "")),
    ]

    for row_index, (label, value) in enumerate(rows, start=3):
        sheet.cell(row=row_index, column=1, value=label).font = Font(bold=True)
        sheet.cell(row=row_index, column=2, value=value)

    start_row = 13
    headers = ["Angle", "Label", "Image URL", "Saved File"]
    for column, header in enumerate(headers, start=1):
        cell = sheet.cell(row=start_row, column=column, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = title_fill

    export_data = _inspection_to_export_dict(inspection)
    for offset, image in enumerate(export_data.get("images", []), start=1):
        row = start_row + offset
        sheet.cell(row=row, column=1, value=image["angle"])
        sheet.cell(row=row, column=2, value=image["label"])
        sheet.cell(row=row, column=3, value=image["url"])
        sheet.cell(row=row, column=4, value=image.get("path", ""))

    workbook.save(report_path)
    return report_path


def _create_photo_report(inspection: Inspection) -> Path:
    inspection_data = _inspection_to_export_dict(inspection)
    report_dir = _inspection_report_dir(inspection.id)
    report_path = report_dir / f"inspection_{inspection.id}_photo_report.docx"

    document = Document()
    document.add_heading("Container Inspection Photo Report", level=1)

    info_table = document.add_table(rows=0, cols=2)
    info_rows = [
        ("Container Number", inspection_data["container_number"]),
        ("Booking Number", inspection_data["booking_number"]),
        ("Truck Number", inspection_data["truck_number"]),
        ("Worker", inspection_data["worker_name"]),
        ("Port / Location", inspection_data["port_name"]),
        ("Status", inspection_data["status"]),
        ("Submitted At", inspection_data["created_at"]),
        ("Notes", inspection_data.get("notes", "")),
    ]

    for label, value in info_rows:
        row = info_table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value

    document.add_heading("Photo Evidence", level=2)

    for image in inspection_data.get("images", []):
        document.add_paragraph(
            f"{image['angle']:02d}. {image['label']}",
            style="Heading 3",
        )
        image_path = _image_path(image)
        if image_path is None:
            document.add_paragraph(image["url"])
            continue

        try:
            document.add_picture(str(image_path), width=Inches(5.8))
        except Exception:
            document.add_paragraph(image["url"])

    document.save(report_path)
    return report_path


def _export_response(
    *,
    request: Request,
    inspection: Inspection,
    report_path: Path,
    export_label: str,
    export_type: str,
    db: Session,
) -> dict[str, Any]:
    inspection_data = _inspection_to_dict(inspection)
    email_sent, email_to, email_message = _email_attachment(
        inspection=inspection_data,
        attachment_path=report_path,
        export_label=export_label,
    )
    report_url = _public_url(request, f"reports/{inspection.id}/{report_path.name}")

    db.add(
        ExportRecord(
            inspection_id=inspection.id,
            export_type=export_type,
            filename=report_path.name,
            report_url=report_url,
            email_sent=email_sent,
            email_to=email_to,
            message=email_message,
        )
    )
    db.commit()

    return {
        "status": "exported",
        "message": f"{export_label} generated. {email_message}",
        "report_url": report_url,
        "email_sent": email_sent,
        "email_to": email_to,
        "filename": report_path.name,
    }


@router.post(
    "/inspections/{inspection_id}/export-excel-email",
    response_model=ExportEmailResponse,
)
async def export_excel_and_email(
    request: Request,
    inspection_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_roles("manager", "admin")),
):
    inspection = _find_inspection(db, inspection_id)
    _ensure_accepted(inspection)
    _ensure_storage()

    report_path = _create_excel_report(inspection)
    return _export_response(
        request=request,
        inspection=inspection,
        report_path=report_path,
        export_label="Excel export",
        export_type="excel",
        db=db,
    )


@router.post(
    "/inspections/{inspection_id}/generate-report-email",
    response_model=ExportEmailResponse,
)
async def generate_report_and_email(
    request: Request,
    inspection_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_roles("manager", "admin")),
):
    inspection = _find_inspection(db, inspection_id)
    _ensure_accepted(inspection)
    _ensure_storage()

    report_path = _create_photo_report(inspection)
    return _export_response(
        request=request,
        inspection=inspection,
        report_path=report_path,
        export_label="Photo report",
        export_type="photo_report",
        db=db,
    )


@router.post(
    "/inspections/{inspection_id}/export-email",
    response_model=ExportEmailResponse,
)
async def export_and_email_inspection(
    request: Request,
    inspection_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_roles("manager", "admin")),
):
    return await export_excel_and_email(request, inspection_id, db, current_user)


@router.post("/ai/scan-container-id", response_model=ScanContainerIdResponse)
async def scan_container_id(
    image: UploadFile = File(...),
    current_user: User = Depends(require_roles("worker", "admin")),
):
    suffix = _validate_image_upload(image)
    _ensure_storage()

    temp_path = OCR_TEMP_ROOT / f"{uuid4()}{suffix}"
    await _save_validated_upload(image, temp_path)

    try:
        container_number = await run_in_threadpool(detect_container_number, temp_path)
    except ContainerOcrError:
        raise HTTPException(
            status_code=422,
            detail="Container number could not be detected from this image",
        )
    finally:
        temp_path.unlink(missing_ok=True)

    return {"container_number": container_number}
