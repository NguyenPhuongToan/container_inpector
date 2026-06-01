from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches
from fastapi import Request

from app.core.constants import REPORT_ROOT
from app.database.models import Inspection


class WordReportError(RuntimeError):
    """Raised when a Word photo report cannot be generated."""


def _image_path(image: dict[str, Any]) -> Path | None:
    path = image.get("path")
    if not path:
        return None

    image_path = Path(path)
    if image_path.exists() and image_path.suffix.lower() in {".jpg", ".jpeg", ".png"}:
        return image_path

    return None


def _inspection_to_export_dict(inspection: Inspection) -> dict[str, Any]:
    images = [
        {
            "angle": image.angle,
            "label": image.label,
            "url": image.url,
            "path": image.path,
        }
        for image in sorted(inspection.images, key=lambda item: item.angle)
    ]
    return {
        "container_number": inspection.container_number,
        "booking_number": inspection.booking_number,
        "truck_number": inspection.truck_number,
        "worker_name": inspection.worker_name,
        "port_name": inspection.port_name,
        "notes": inspection.notes or "",
        "status": inspection.status,
        "created_at": inspection.created_at.isoformat(),
        "images": images,
    }


def _report_dir(inspection: Inspection) -> Path:
    report_dir = REPORT_ROOT / inspection.id
    report_dir.mkdir(parents=True, exist_ok=True)
    return report_dir


def generate_word_report(inspection: Inspection, request: Request) -> Path:
    try:
        inspection_data = _inspection_to_export_dict(inspection)
        report_path = _report_dir(inspection) / f"inspection_{inspection.id}_photo_report.docx"

        document = Document()
        document.add_heading("Container Inspection Photo Report", level=1)

        info_table = document.add_table(rows=0, cols=2)
        info_rows = [
            ("Container Number", inspection_data["container_number"]),
            ("Booking Number", inspection_data["booking_number"]),
            ("Truck Number", inspection_data["truck_number"]),
            ("Worker", inspection_data["worker_name"]),
            ("Port / Location", inspection_data["port_name"]),
            ("Status", inspection_data["status"]),
            ("Submitted At", inspection_data["created_at"]),
            ("Notes", inspection_data.get("notes", "")),
        ]

        for label, value in info_rows:
            row = info_table.add_row()
            row.cells[0].text = label
            row.cells[1].text = value

        document.add_heading("Photo Evidence", level=2)

        for image in inspection_data.get("images", []):
            document.add_paragraph(
                f"{image['angle']:02d}. {image['label']}",
                style="Heading 3",
            )
            image_path = _image_path(image)
            if image_path is None:
                document.add_paragraph(image["url"])
                continue

            try:
                document.add_picture(str(image_path), width=Inches(5.8))
            except Exception:
                document.add_paragraph(image["url"])

        document.save(report_path)
        return report_path
    except Exception as exc:
        raise WordReportError("Could not generate Word photo report") from exc
